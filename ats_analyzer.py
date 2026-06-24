import streamlit as st
import requests
import json
import re
import os
import io
import time
import fitz
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from rapidfuzz import fuzz, process

OPENROUTER_API_KEY = os.environ.get("OPENROUTER_API_KEY", "sk-or-v1-????") # APİ KEY
OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"
MODEL = "openrouter/free"
RATE_LIMIT_DELAY = 8

st.set_page_config(
    page_title="CVMatch",
    page_icon="⭐",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("<style>#MainMenu{visibility:hidden;}[data-testid='stToolbar']{display:none;}.stDeployButton{display:none;}header[data-testid='stHeader'] button{display:none;}:root{--primary:#f97316;--primary2:#fb923c;--success:#22c55e;--warning:#f59e0b;--danger:#ef4444;}.stApp{background:radial-gradient(circle at top left,#431407 0%,transparent 35%),radial-gradient(circle at top right,#7c2d12 0%,transparent 30%),linear-gradient(180deg,#0f172a 0%,#020617 100%);}.block-container{max-width:1400px;padding-top:2rem;padding-bottom:2rem;}.main-header{background:rgba(255,255,255,.05);backdrop-filter:blur(16px);border:1px solid rgba(249,115,22,.15);border-radius:18px;padding:18px 24px;text-align:center;color:white;margin-bottom:20px;box-shadow:0 8px 24px rgba(0,0,0,.18);}.main-header h1{font-size:1.9rem;font-weight:700;margin:0 0 6px 0;background:linear-gradient(90deg,#ffffff,#fdba74);-webkit-background-clip:text;-webkit-text-fill-color:transparent;}.main-header p{font-size:0.95rem;color:#cbd5e1;margin:0;}.metric-card{background:rgba(255,255,255,.05);backdrop-filter:blur(16px);border:1px solid rgba(249,115,22,.12);border-radius:20px;padding:24px;transition:.25s ease;box-shadow:0 8px 30px rgba(0,0,0,.20);}.metric-card:hover{transform:translateY(-4px);border-color:rgba(249,115,22,.45);box-shadow:0 15px 40px rgba(249,115,22,.15);}.score-green,.score-yellow,.score-red{font-size:3rem;font-weight:800;letter-spacing:-1px;}.score-green{color:var(--success);}.score-yellow{color:var(--warning);}.score-red{color:var(--danger);}.section-header{color:white;font-size:1.4rem;font-weight:700;margin-bottom:20px;padding-bottom:10px;border-bottom:1px solid rgba(249,115,22,.30);}.skill-badge-match{display:inline-block;padding:8px 14px;margin:4px;border-radius:999px;background:rgba(34,197,94,.15);color:#86efac;border:1px solid rgba(34,197,94,.25);font-size:.85rem;font-weight:600;}.skill-badge-missing{display:inline-block;padding:8px 14px;margin:4px;border-radius:999px;background:rgba(239,68,68,.12);color:#fca5a5;border:1px solid rgba(239,68,68,.22);font-size:.85rem;font-weight:600;}.stButton>button{width:100%;border:none;border-radius:14px;padding:14px;color:white;font-weight:700;background:linear-gradient(135deg,#f97316,#fb923c);transition:.25s ease;box-shadow:0 10px 25px rgba(249,115,22,.30);}.stButton>button:hover{transform:translateY(-2px);box-shadow:0 15px 35px rgba(249,115,22,.45);}.stTextInput input,.stTextArea textarea{background:rgba(255,255,255,.04)!important;border:1px solid rgba(249,115,22,.15)!important;color:white!important;border-radius:12px!important;}.stTextInput input:focus,.stTextArea textarea:focus{border-color:#f97316!important;box-shadow:0 0 0 3px rgba(249,115,22,.20)!important;}[data-testid='stFileUploader']{background:rgba(255,255,255,.04);border:1px dashed rgba(249,115,22,.25);border-radius:18px;padding:18px;}[data-testid='stDataFrame']{border-radius:18px;overflow:hidden;border:1px solid rgba(249,115,22,.12);}[data-testid='metric-container']{background:rgba(255,255,255,.04);border:1px solid rgba(249,115,22,.12);border-radius:18px;padding:15px;}::-webkit-scrollbar{width:10px;}::-webkit-scrollbar-track{background:#0f172a;}::-webkit-scrollbar-thumb{background:#f97316;border-radius:999px;}@media(max-width:768px){.main-header h1{font-size:2rem;}.score-green,.score-yellow,.score-red{font-size:2.2rem;}}</style>", unsafe_allow_html=True)


def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    text = ""
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        for page in doc:
            text += page.get_text()
        doc.close()
    except Exception as e:
        st.error(f"PDF okuma hatası: {e}")
    return text.strip()


def call_openrouter(prompt: str, system_prompt: str = "") -> str:
    active_key = os.environ.get("OPENROUTER_API_KEY", OPENROUTER_API_KEY)
    if not active_key:
        st.error("OPENROUTER_API_KEY bulunamadı. Sol panelden API key'inizi girin.")
        return ""
    headers = {
        "Authorization": f"Bearer {active_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://cvmatch.app",
        "X-Title": "CVMatch"
    }
    messages = []
    if system_prompt:
        messages.append({"role": "system", "content": system_prompt})
    messages.append({"role": "user", "content": prompt})
    payload = {
        "model": MODEL,
        "messages": messages,
        "temperature": 0.3,
        "max_tokens": 4000
    }
    max_retries = 3
    for attempt in range(max_retries):
        try:
            response = requests.post(OPENROUTER_URL, headers=headers, json=payload, timeout=60)
            if response.status_code == 429:
                wait_time = RATE_LIMIT_DELAY * (attempt + 1)
                st.warning(f"⏳ Rate limit aşıldı, {wait_time} saniye bekleniyor... (Deneme {attempt + 1}/{max_retries})")
                time.sleep(wait_time)
                continue
            response.raise_for_status()
            data = response.json()
            time.sleep(RATE_LIMIT_DELAY)
            return data["choices"][0]["message"]["content"]
        except requests.exceptions.Timeout:
            st.error("API isteği zaman aşımına uğradı. Lütfen tekrar deneyin.")
            return ""
        except requests.exceptions.HTTPError as e:
            if response.status_code == 429 and attempt < max_retries - 1:
                wait_time = RATE_LIMIT_DELAY * (attempt + 2)
                st.warning(f"⏳ Rate limit, {wait_time}s bekleniyor...")
                time.sleep(wait_time)
                continue
            st.error(f"API bağlantı hatası: {e}")
            return ""
        except requests.exceptions.RequestException as e:
            st.error(f"API bağlantı hatası: {e}")
            return ""
        except (KeyError, IndexError) as e:
            st.error(f"API yanıt ayrıştırma hatası: {e}")
            return ""
    st.error("Maksimum deneme sayısına ulaşıldı. Lütfen birkaç dakika bekleyip tekrar deneyin.")
    return ""


def parse_json_from_response(response: str) -> dict:
    try:
        json_match = re.search(r'\{.*\}', response, re.DOTALL)
        if json_match:
            return json.loads(json_match.group())
    except json.JSONDecodeError:
        pass
    return {}


def extract_cv_data(cv_text: str) -> dict:
    system_prompt = "Sen bir CV analiz uzmanısın. Sadece JSON formatında yanıt ver."
    prompt = f"""Aşağıdaki CV metnini analiz et ve şu JSON formatında döndür:
{{
  "skills": ["skill1", "skill2"],
  "experience_years": 0,
  "education": "eğitim bilgisi",
  "languages": ["dil1", "dil2"],
  "job_titles": ["ünvan1"],
  "summary": "kısa özet"
}}

CV metni:
{cv_text[:3000]}
"""
    response = call_openrouter(prompt, system_prompt)
    result = parse_json_from_response(response)
    defaults = {
        "skills": [],
        "experience_years": 0,
        "education": "",
        "languages": [],
        "job_titles": [],
        "summary": ""
    }
    defaults.update(result)
    return defaults


def extract_job_data(job_text: str) -> dict:
    system_prompt = "Sen bir iş ilanı analiz uzmanısın. Sadece JSON formatında yanıt ver."
    prompt = f"""Aşağıdaki iş ilanını analiz et ve şu JSON formatında döndür:
{{
  "required_skills": ["skill1", "skill2"],
  "preferred_skills": ["skill3"],
  "min_experience_years": 0,
  "education_requirement": "lisans",
  "job_title": "pozisyon adı",
  "key_responsibilities": ["görev1", "görev2"]
}}

İş ilanı:
{job_text[:3000]}
"""
    response = call_openrouter(prompt, system_prompt)
    result = parse_json_from_response(response)
    defaults = {
        "required_skills": [],
        "preferred_skills": [],
        "min_experience_years": 0,
        "education_requirement": "",
        "job_title": "",
        "key_responsibilities": []
    }
    defaults.update(result)
    return defaults


def calculate_skill_match(cv_skills: list, job_skills: list) -> tuple[float, list, list]:
    if not job_skills:
        return 0.0, [], []
    matched = []
    unmatched = []
    cv_skills_lower = [s.lower().strip() for s in cv_skills]
    for job_skill in job_skills:
        job_skill_lower = job_skill.lower().strip()
        best_match = process.extractOne(job_skill_lower, cv_skills_lower, scorer=fuzz.token_set_ratio)
        if best_match and best_match[1] >= 75:
            matched.append(job_skill)
        else:
            unmatched.append(job_skill)
    score = (len(matched) / len(job_skills)) * 100 if job_skills else 0
    return round(score, 1), matched, unmatched


def calculate_ats_score(cv_data: dict, job_data: dict) -> dict:
    required_skills = job_data.get("required_skills", [])
    preferred_skills = job_data.get("preferred_skills", [])
    cv_skills = cv_data.get("skills", [])

    req_score, req_matched, req_missing = calculate_skill_match(cv_skills, required_skills)
    pref_score, pref_matched, _ = calculate_skill_match(cv_skills, preferred_skills)

    all_matched = list(set(req_matched + pref_matched))
    all_missing = list(set(req_missing))

    cv_exp = cv_data.get("experience_years", 0)
    min_exp = job_data.get("min_experience_years", 0)
    if min_exp == 0:
        exp_score = 100.0
    elif cv_exp >= min_exp:
        exp_score = 100.0
    elif cv_exp >= min_exp * 0.7:
        exp_score = 70.0
    else:
        exp_score = max(0.0, (cv_exp / max(min_exp, 1)) * 100)

    cv_edu = cv_data.get("education", "").lower()
    job_edu = job_data.get("education_requirement", "").lower()
    edu_score = 100.0
    if job_edu:
        edu_keywords = ["doktora", "phd", "yüksek lisans", "master", "lisans", "bachelor", "önlisans", "associate"]
        job_edu_level = 0
        cv_edu_level = 0
        for i, kw in enumerate(edu_keywords):
            if kw in job_edu:
                job_edu_level = max(job_edu_level, len(edu_keywords) - i)
            if kw in cv_edu:
                cv_edu_level = max(cv_edu_level, len(edu_keywords) - i)
        if cv_edu_level >= job_edu_level:
            edu_score = 100.0
        elif cv_edu_level > 0:
            edu_score = 65.0
        else:
            edu_score = 40.0

    total_score = (req_score * 0.55) + (pref_score * 0.15) + (exp_score * 0.20) + (edu_score * 0.10)

    return {
        "total_score": round(total_score, 1),
        "skill_match_score": round(req_score, 1),
        "preferred_skill_score": round(pref_score, 1),
        "experience_score": round(exp_score, 1),
        "education_score": round(edu_score, 1),
        "matched_skills": all_matched,
        "missing_skills": all_missing,
        "required_matched": req_matched,
        "required_missing": req_missing
    }


def find_missing_skills(missing_skills: list, job_text: str) -> list:
    if not missing_skills:
        return []
    system_prompt = "Sen bir kariyer koçusun. Sadece JSON formatında yanıt ver."
    prompt = f"""Bir CV'de şu beceriler eksik: {', '.join(missing_skills)}

İş ilanı: {job_text[:1000]}

Her eksik beceri için öğrenme önerileri ver:
{{
  "missing_skills_analysis": [
    {{
      "skill": "beceri adı",
      "importance": "yüksek/orta/düşük",
      "learn_suggestion": "nasıl öğrenilir",
      "time_to_learn": "yaklaşık süre"
    }}
  ]
}}"""
    response = call_openrouter(prompt, system_prompt)
    result = parse_json_from_response(response)
    return result.get("missing_skills_analysis", [])


def optimize_cv(cv_text: str, job_text: str, missing_skills: list, ats_score: dict) -> str:
    system_prompt = "Sen profesyonel bir CV yazarısın. ATS optimizasyonu konusunda uzmansın."
    prompt = f"""Aşağıdaki CV'yi, verilen iş ilanına göre ATS uyumlu şekilde optimize et.

Mevcut ATS skoru: {ats_score.get('total_score', 0)}/100
Eksik beceriler: {', '.join(missing_skills[:10])}

Yaptığın işlemler:
1. Eksik anahtar kelimeleri uygun yerlere ekle
2. Profesyonel özet bölümü oluştur veya güçlendir
3. Teknik becerileri iş ilanına göre öne çıkar
4. ATS sistemlerinin anlayacağı net başlıklar kullan
5. İş ilanındaki önemli terimleri CV'de kullan

CV:
{cv_text[:2500]}

İş İlanı:
{job_text[:1500]}

Optimize edilmiş CV'yi Markdown formatında yaz. Bölüm başlıklarını ## ile yaz."""
    response = call_openrouter(prompt, system_prompt)
    return response if response else cv_text


def generate_docx(optimized_cv_text: str, candidate_name: str = "Aday") -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(candidate_name)
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x46, 0x6F, 0xEA)

    lines = optimized_cv_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        if line.startswith('## '):
            heading_para = doc.add_paragraph()
            heading_run = heading_para.add_run(line[3:])
            heading_run.font.size = Pt(13)
            heading_run.font.bold = True
            heading_run.font.color.rgb = RGBColor(0x46, 0x6F, 0xEA)
            heading_para.paragraph_format.space_before = Pt(10)
            heading_para.paragraph_format.space_after = Pt(4)
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            pPr = heading_para._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '4670EA')
            pBdr.append(bottom)
            pPr.append(pBdr)

        elif line.startswith('# '):
            heading_para = doc.add_paragraph()
            heading_run = heading_para.add_run(line[2:])
            heading_run.font.size = Pt(15)
            heading_run.font.bold = True
            heading_para.paragraph_format.space_before = Pt(12)

        elif line.startswith('- ') or line.startswith('* '):
            bullet_para = doc.add_paragraph(style='List Bullet')
            bullet_para.add_run(line[2:]).font.size = Pt(10)

        elif line.startswith('**') and line.endswith('**'):
            bold_para = doc.add_paragraph()
            bold_run = bold_para.add_run(line[2:-2])
            bold_run.font.bold = True
            bold_run.font.size = Pt(10)

        else:
            clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            clean_line = re.sub(r'\*(.*?)\*', r'\1', clean_line)
            para = doc.add_paragraph()
            para.add_run(clean_line).font.size = Pt(10)

    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("Bu CV CVMatch ile optimize edilmiştir.")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def render_score_color(score: float) -> str:
    if score >= 80:
        return "score-green"
    elif score >= 60:
        return "score-yellow"
    else:
        return "score-red"


def render_dashboard(ats_result: dict, cv_data: dict, job_data: dict) -> None:
    total = ats_result.get("total_score", 0)
    skill_score = ats_result.get("skill_match_score", 0)
    exp_score = ats_result.get("experience_score", 0)
    edu_score = ats_result.get("education_score", 0)
    matched = ats_result.get("matched_skills", [])
    missing = ats_result.get("missing_skills", [])

    score_class = render_score_color(total)

    st.markdown('<h2 class="section-header">📊 CVMatch Analiz Sonuçları</h2>', unsafe_allow_html=True)

    col1, col2, col3, col4 = st.columns(4)

    with col1:
        st.markdown(f"""
        <div class="metric-card">
            <div style="font-size:0.9rem; color:#666; margin-bottom:8px;">⭐ ATS Toplam Skoru</div>
            <div class="{score_class}">{total}/100</div>
        </div>
        """, unsafe_allow_html=True)

    with col2:
        sc = render_score_color(skill_score)
        st.markdown(f"""
        <div class="metric-card">
            <div style="font-size:0.9rem; color:#666; margin-bottom:8px;">🛠️ Beceri Eşleşmesi</div>
            <div class="{sc}">{skill_score}%</div>
        </div>
        """, unsafe_allow_html=True)

    with col3:
        ec = render_score_color(exp_score)
        st.markdown(f"""
        <div class="metric-card">
            <div style="font-size:0.9rem; color:#666; margin-bottom:8px;">📅 Deneyim Skoru</div>
            <div class="{ec}">{exp_score}%</div>
        </div>
        """, unsafe_allow_html=True)

    with col4:
        edc = render_score_color(edu_score)
        st.markdown(f"""
        <div class="metric-card">
            <div style="font-size:0.9rem; color:#666; margin-bottom:8px;">🎓 Eğitim Skoru</div>
            <div class="{edc}">{edu_score}%</div>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    col_left, col_right = st.columns(2)

    with col_left:
        st.markdown('<h3 class="section-header">✅ Eşleşen Beceriler</h3>', unsafe_allow_html=True)
        if matched:
            badges_html = "".join([f'<span class="skill-badge-match">{s}</span>' for s in matched])
            st.markdown(f'<div>{badges_html}</div>', unsafe_allow_html=True)
        else:
            st.info("Eşleşen beceri bulunamadı.")

    with col_right:
        st.markdown('<h3 class="section-header">❌ Eksik Beceriler</h3>', unsafe_allow_html=True)
        if missing:
            badges_html = "".join([f'<span class="skill-badge-missing">{s}</span>' for s in missing])
            st.markdown(f'<div>{badges_html}</div>', unsafe_allow_html=True)
        else:
            st.success("Tüm gerekli beceriler mevcut!")

    st.markdown("<br>", unsafe_allow_html=True)

    score_data = {
        "Kategori": ["Zorunlu Beceri Eşleşmesi", "Tercih Edilen Beceriler", "Deneyim", "Eğitim"],
        "Skor": [
            ats_result.get("skill_match_score", 0),
            ats_result.get("preferred_skill_score", 0),
            exp_score,
            edu_score
        ],
        "Ağırlık": ["55%", "15%", "20%", "10%"]
    }
    df = pd.DataFrame(score_data)
    st.markdown('<h3 class="section-header">📈 Detaylı Skor Analizi</h3>', unsafe_allow_html=True)
    st.dataframe(df, use_container_width=True, hide_index=True)

    if cv_data.get("skills"):
        st.markdown('<h3 class="section-header">👤 CV Özeti</h3>', unsafe_allow_html=True)
        info_col1, info_col2 = st.columns(2)
        with info_col1:
            st.metric("Toplam Deneyim", f"{cv_data.get('experience_years', 0)} yıl")
            st.metric("Eğitim", cv_data.get("education", "Belirtilmemiş")[:50])
        with info_col2:
            st.metric("Tespit Edilen Beceri Sayısı", len(cv_data.get("skills", [])))
            languages = cv_data.get("languages", [])
            st.metric("Bildiği Diller", ", ".join(languages) if languages else "Belirtilmemiş")


def main() -> None:
    st.markdown("""
    <div class="main-header">
        <h1>⭐ CVMatch</h1>
        <p>Akıllı CV ve İş İlanı Eşleştirme Sistemi</p>
    </div>
    """, unsafe_allow_html=True)

    with st.sidebar:
        st.markdown("### ⚙️ Ayarlar")
        api_key_input = st.text_input(
            "OpenRouter API Key",
            value=OPENROUTER_API_KEY,
            type="password",
            help="OPENROUTER_API_KEY env variable veya buradan girin"
        )
        if api_key_input:
            os.environ["OPENROUTER_API_KEY"] = api_key_input
            if api_key_input != OPENROUTER_API_KEY:
                st.success("API Key güncellendi!")

        st.markdown("---")
        st.markdown("### 📋 Nasıl Kullanılır?")
        st.markdown("""
        1. CV'nizi PDF olarak yükleyin
        2. İş ilanı metnini yapıştırın
        3. **Analiz Et** butonuna basın
        4. Sonuçları inceleyin
        5. Optimize edilmiş CV'yi indirin
        """)

        st.markdown("---")
        st.markdown("### 🤖 Model")
        st.info(f"`{MODEL}`")

    tab1, tab2, tab3, tab4 = st.tabs([
        "📤 CV Yükle & Analiz",
        "📊 Analiz Sonuçları",
        "✨ CV Optimizasyonu",
        "📥 Export"
    ])

    if "analysis_done" not in st.session_state:
        st.session_state.analysis_done = False
    if "cv_data" not in st.session_state:
        st.session_state.cv_data = {}
    if "job_data" not in st.session_state:
        st.session_state.job_data = {}
    if "ats_result" not in st.session_state:
        st.session_state.ats_result = {}
    if "optimized_cv" not in st.session_state:
        st.session_state.optimized_cv = ""
    if "cv_text" not in st.session_state:
        st.session_state.cv_text = ""
    if "job_text" not in st.session_state:
        st.session_state.job_text = ""
    if "missing_analysis" not in st.session_state:
        st.session_state.missing_analysis = []

    with tab1:
        st.markdown('<h2 class="section-header">📤 CV ve İş İlanı Girişi</h2>', unsafe_allow_html=True)

        upload_col, job_col = st.columns(2)

        with upload_col:
            st.markdown("#### 📄 CV Yükle")
            uploaded_file = st.file_uploader(
                "PDF formatında CV yükleyin",
                type=["pdf"],
                help="Maksimum dosya boyutu: 10MB"
            )
            if uploaded_file:
                st.success(f"✅ {uploaded_file.name} yüklendi ({uploaded_file.size // 1024} KB)")
                cv_bytes = uploaded_file.read()
                cv_text_preview = extract_text_from_pdf(cv_bytes)
                if cv_text_preview:
                    st.session_state.cv_text = cv_text_preview
                    with st.expander("CV Önizleme (ilk 500 karakter)"):
                        st.text(cv_text_preview[:500] + "...")

        with job_col:
            st.markdown("#### 💼 İş İlanı")
            job_text_input = st.text_area(
                "İş ilanı metnini yapıştırın",
                height=300,
                placeholder="Buraya iş ilanı metnini yapıştırın...\n\nÖrnek:\nPozisyon: Senior Python Developer\nGereksinimler:\n- Python 5+ yıl\n- Docker, Kubernetes\n- AWS veya Azure\n..."
            )
            if job_text_input:
                st.session_state.job_text = job_text_input

        st.markdown("<br>", unsafe_allow_html=True)

        analyze_btn = st.button("🚀 Analiz Et", use_container_width=True)

        if analyze_btn:
            current_api_key = os.environ.get("OPENROUTER_API_KEY", OPENROUTER_API_KEY)
            if not current_api_key:
                st.error("❌ OpenRouter API Key girilmedi. Sol panelden API key'inizi girin.")
            elif not st.session_state.cv_text:
                st.error("❌ Lütfen bir CV yükleyin.")
            elif not st.session_state.job_text:
                st.error("❌ Lütfen iş ilanı metnini girin.")
            else:
                progress_bar = st.progress(0)
                status_text = st.empty()

                try:
                    status_text.text("🔍 CV analiz ediliyor...")
                    progress_bar.progress(20)
                    cv_data = extract_cv_data(st.session_state.cv_text)
                    st.session_state.cv_data = cv_data

                    status_text.text("📋 İş ilanı analiz ediliyor...")
                    progress_bar.progress(40)
                    job_data = extract_job_data(st.session_state.job_text)
                    st.session_state.job_data = job_data

                    status_text.text("⭐ ATS skoru hesaplanıyor...")
                    progress_bar.progress(60)
                    ats_result = calculate_ats_score(cv_data, job_data)
                    st.session_state.ats_result = ats_result

                    status_text.text("🔎 Eksik beceriler analiz ediliyor...")
                    progress_bar.progress(75)
                    if ats_result.get("missing_skills"):
                        missing_analysis = find_missing_skills(
                            ats_result["missing_skills"][:8],
                            st.session_state.job_text
                        )
                        st.session_state.missing_analysis = missing_analysis

                    st.session_state.analysis_done = True
                    progress_bar.progress(100)
                    status_text.text("✅ Analiz tamamlandı!")

                    total_score = ats_result.get("total_score", 0)
                    if total_score >= 80:
                        st.success(f"🎉 Harika! Uyum Skoru: {total_score}/100 - CV'niz bu pozisyon için çok uygun!")
                    elif total_score >= 60:
                        st.warning(f"⚠️ Uyum Skoru: {total_score}/100 - CV'niz optimize edilebilir.")
                    else:
                        st.error(f"❌ Uyum Skoru: {total_score}/100 - CV'niz optimize edilmeli.")

                    st.info("📊 Detaylar için 'Analiz Sonuçları' sekmesine geçin.")

                except Exception as e:
                    st.error(f"Analiz sırasında hata oluştu: {e}")
                    progress_bar.progress(0)

    with tab2:
        if not st.session_state.analysis_done:
            st.info("⏳ Henüz analiz yapılmadı. 'CV Yükle & Analiz' sekmesinden analizi başlatın.")
        else:
            render_dashboard(
                st.session_state.ats_result,
                st.session_state.cv_data,
                st.session_state.job_data
            )

            if st.session_state.missing_analysis:
                st.markdown('<h3 class="section-header">📚 Eksik Beceri Geliştirme Önerileri</h3>', unsafe_allow_html=True)
                for item in st.session_state.missing_analysis:
                    importance = item.get("importance", "orta")
                    icon = "🔴" if importance == "yüksek" else "🟡" if importance == "orta" else "🟢"
                    with st.expander(f"{icon} {item.get('skill', 'Beceri')} ({importance} öncelik)"):
                        col_a, col_b = st.columns(2)
                        with col_a:
                            st.markdown(f"**Öğrenme Yolu:** {item.get('learn_suggestion', 'N/A')}")
                        with col_b:
                            st.markdown(f"**Tahmini Süre:** {item.get('time_to_learn', 'N/A')}")

    with tab3:
        if not st.session_state.analysis_done:
            st.info("⏳ Önce analiz yapmanız gerekiyor.")
        else:
            st.markdown('<h2 class="section-header">✨ CV Optimizasyonu</h2>', unsafe_allow_html=True)

            total_score = st.session_state.ats_result.get("total_score", 0)
            missing_count = len(st.session_state.ats_result.get("missing_skills", []))

            info_cols = st.columns(3)
            with info_cols[0]:
                st.metric("Mevcut Uyum Skoru", f"{total_score}/100")
            with info_cols[1]:
                st.metric("Eksik Beceri Sayısı", missing_count)
            with info_cols[2]:
                st.metric("Tahmini Yeni Skor", f"{min(total_score + 15, 98)}/100", delta="+15")

            st.markdown("<br>", unsafe_allow_html=True)
            optimize_btn = st.button("🤖 CV'yi Optimize Et", use_container_width=True)

            if optimize_btn:
                with st.spinner("🔄 CV optimize ediliyor... Bu işlem 30-60 saniye sürebilir."):
                    try:
                        optimized = optimize_cv(
                            st.session_state.cv_text,
                            st.session_state.job_text,
                            st.session_state.ats_result.get("missing_skills", []),
                            st.session_state.ats_result
                        )
                        st.session_state.optimized_cv = optimized
                        st.success("✅ CV başarıyla optimize edildi!")
                    except Exception as e:
                        st.error(f"Optimizasyon hatası: {e}")

            if st.session_state.optimized_cv:
                st.markdown('<h3 class="section-header">📄 Optimize Edilmiş CV</h3>', unsafe_allow_html=True)

                copy_col, _ = st.columns([1, 3])
                with copy_col:
                    if st.button("📋 Metni Kopyala"):
                        st.code(st.session_state.optimized_cv)
                        st.info("Yukarıdaki metni kopyalayabilirsiniz.")

                st.markdown(st.session_state.optimized_cv)

                st.markdown("---")
                st.info("📥 DOCX olarak indirmek için 'Export' sekmesine geçin.")

    with tab4:
        st.markdown('<h2 class="section-header">📥 CV Export</h2>', unsafe_allow_html=True)

        if not st.session_state.optimized_cv:
            st.warning("⚠️ Henüz optimize edilmiş CV yok. 'CV Optimizasyonu' sekmesinden optimize edin.")

            if st.session_state.cv_text:
                st.markdown("#### Ham CV'yi İndir")
                st.info("Orijinal CV'nizi de DOCX olarak indirebilirsiniz.")
                if st.button("📄 Orijinal CV'yi DOCX İndir"):
                    try:
                        docx_bytes = generate_docx(st.session_state.cv_text, "CV")
                        st.download_button(
                            label="⬇️ DOCX İndir",
                            data=docx_bytes,
                            file_name="original_cv.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    except Exception as e:
                        st.error(f"DOCX oluşturma hatası: {e}")
        else:
            st.markdown("#### 📊 Export Özeti")
            export_cols = st.columns(3)
            with export_cols[0]:
                st.metric("Uyum Skoru", f"{st.session_state.ats_result.get('total_score', 0)}/100")
            with export_cols[1]:
                st.metric("Eşleşen Beceri", len(st.session_state.ats_result.get("matched_skills", [])))
            with export_cols[2]:
                st.metric("Optimize Edildi", "✅ Evet")

            st.markdown("<br>", unsafe_allow_html=True)

            candidate_name = st.text_input(
                "Adınız (DOCX için)",
                placeholder="Adınız Soyadınız",
                help="DOCX dosyasının başlığında kullanılacak"
            )

            dl_col1, dl_col2 = st.columns(2)

            with dl_col1:
                if st.button("🔄 DOCX Oluştur", use_container_width=True):
                    try:
                        name = candidate_name if candidate_name else "Aday"
                        docx_bytes = generate_docx(st.session_state.optimized_cv, name)
                        st.session_state.docx_bytes = docx_bytes
                        st.success("✅ DOCX hazır! İndir butonuna basın.")
                    except Exception as e:
                        st.error(f"DOCX oluşturma hatası: {e}")

            with dl_col2:
                if "docx_bytes" in st.session_state and st.session_state.docx_bytes:
                    name = candidate_name if candidate_name else "optimize_cv"
                    safe_name = re.sub(r'[^\w\s-]', '', name).strip().replace(' ', '_')
                    filename = f"{safe_name}_CVMatch_Optimized.docx"
                    st.download_button(
                        label="⬇️ DOCX İndir",
                        data=st.session_state.docx_bytes,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )

            st.markdown("---")
            st.markdown("#### 📋 Metin Olarak Kopyala")
            st.text_area(
                "Optimize edilmiş CV metni",
                value=st.session_state.optimized_cv,
                height=400,
                help="Bu metni kopyalayabilirsiniz"
            )

    st.markdown("---")
    st.markdown(
        "<div style='text-align:center; color:#999; font-size:0.8rem;'>"
        "<a href='https://github.com/denzworks' target='_blank' style='color:#667eea; text-decoration:none;'>GitHub</a>"
        " | "
        "CVMatch v1.0 | "
        "<a href='https://openrouter.ai/keys' target='_blank' style='color:#667eea; text-decoration:none;'>Api Key</a>"
        "</div>",
        unsafe_allow_html=True
    )


if __name__ == "__main__":
    main()
